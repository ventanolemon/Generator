from PyQt6 import uic
from PyQt6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QPushButton, QComboBox, QFileDialog, \
    QMessageBox, QMainWindow
from docx import Document
from docx.shared import Pt
import os


class ExportNewDialog(QMainWindow):
    def __init__(self, test_name, answer_visibility, parent=None):
        super().__init__(parent)
        self.test_text = None
        self.answers = None
        self.test_name = test_name
        self.answer_visibility = answer_visibility

        self.init_ui()
        self.setup_connections()

    def init_ui(self):
        uic.loadUi('resources/exporter/export_new/export_new_variants.ui', self)
        self.fileName.setText(self.test_name)
        self.variantCount.setValue(1)

    def setup_connections(self):
        self.chooseDirectoryButton.clicked.connect(self.choose_directory)
        self.exportButton.clicked.connect(self.export_to_word)

    def choose_directory(self):
        directory = QFileDialog.getExistingDirectory(
            self,
            "Выберите директорию",
            os.path.expanduser("~"),
            QFileDialog.Option.ShowDirsOnly
        )
        if directory:
            self.pathDirectory.setText(directory)

    def export_to_word(self):
        try:
            # Валидация данных
            if not self.pathDirectory.text():
                QMessageBox.warning(self, "Ошибка", "Выберите директорию для сохранения!")
                return

            if not self.fileName.text().strip():
                QMessageBox.warning(self, "Ошибка", "Введите название файла!")
                return

            # Получаем количество вариантов
            variant_count = self.variantCount.value()
            try:
                variant_count = int(variant_count)
                if variant_count < 1:
                    raise ValueError
            except ValueError:
                QMessageBox.warning(self, "Ошибка", "Некорректное количество вариантов!")
                return

            # Создаем документ
            doc = Document()
            self.configure_styles(doc)

            # Получаем настройку отображения ответов
            answer_option = self.selectAnswers.currentText()

            # Генерируем варианты
            for variant in range(1, variant_count + 1):
                # Добавляем раздел для варианта
                doc.add_heading(f"Вариант {variant}", level=1)

                # Получаем актуальные данные через родительский класс
                test_data = self.parent().get_test()
                if not test_data or len(test_data) != 2:
                    QMessageBox.warning(self, "Ошибка", "Не удалось получить данные теста!")
                    return

                test_text, test_answers = test_data
                content = []

                # Формируем контент в зависимости от настроек
                if answer_option == "показывать рядом с заданием":
                    content = [f"{t}\nОтвет: {a.split(' ', 1)[1]}"
                               for t, a in zip(test_text, test_answers)]
                elif answer_option == "показывать в конце":
                    content = test_text.copy()
                    content.append("\nОтветы:")
                    content.extend(test_answers)
                else:
                    content = test_text

                # Добавляем контент в документ
                for paragraph in content:
                    if isinstance(paragraph, str):
                        if paragraph.startswith("Ответ:"):
                            p = doc.add_paragraph()
                            run = p.add_run(paragraph)
                            run.italic = True
                        else:
                            doc.add_paragraph(paragraph)

                # Добавляем разрыв страницы между вариантами
                if variant < variant_count:
                    doc.add_page_break()

            # Сохраняем документ
            file_name = f"{self.fileName.text().strip()}.docx"
            full_path = os.path.join(self.pathDirectory.text(), file_name)

            # Проверка и создание директории
            os.makedirs(os.path.dirname(full_path), exist_ok=True)

            try:
                doc.save(full_path)
                QMessageBox.information(self, "Успех", f"Файл сохранен:\n{full_path}")
                self.hide()
            except PermissionError:
                QMessageBox.critical(self, "Ошибка", "Нет прав для записи в выбранную директорию!")
            except Exception as e:
                raise e

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка при экспорте: {str(e)}")
            raise e

    def configure_styles(self, doc):
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        heading_style = doc.styles['Heading 1']
        heading_font = heading_style.font
        heading_font.bold = True
        heading_font.size = Pt(16)

