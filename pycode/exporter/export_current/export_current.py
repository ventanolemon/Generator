from PyQt6 import uic
from PyQt6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QPushButton, QComboBox, QFileDialog, \
    QMessageBox, QMainWindow
from docx import Document
from docx.shared import Pt
import os


class ExportDialog(QMainWindow):
    def __init__(self, test_text, answers, test_name, answer_visibility, parent=None):
        super().__init__(parent)
        self.test_text = test_text
        self.answers = answers
        self.test_name = test_name
        self.answer_visibility = answer_visibility

        self.init_ui()
        self.setup_connections()

    def init_ui(self):
        uic.loadUi('pycode/exporter/export_current/export_current_interface.ui', self)
        self.fileName.setText(self.test_name)
        # self.setWindowTitle("Экспорт теста")
        # self.setMinimumWidth(400)
        #
        # layout = QVBoxLayout()
        #
        # # Поле для названия файла
        # self.fileName = QLineEdit(self.test_name)
        # layout.addWidget(QLabel("Название файла:"))
        # layout.addWidget(self.fileName)
        #
        # # Выбор директории
        # dir_layout = QHBoxLayout()
        # self.pathDirectory = QLineEdit()
        # self.pathDirectory.setReadOnly(True)
        # self.chooseDirectoryButton = QPushButton("Выбрать папку...")
        # dir_layout.addWidget(self.pathDirectory)
        # dir_layout.addWidget(self.chooseDirectoryButton)
        # layout.addWidget(QLabel("Директория для сохранения:"))
        # layout.addLayout(dir_layout)
        #
        # # Настройки отображения ответов
        # self.selectAnswers = QComboBox()
        # self.selectAnswers.addItems([
        #     "показывать рядом с заданием",
        #     "показывать в конце",
        #     "скрыть"
        # ])
        # self.selectAnswers.setCurrentText(self.answer_visibility)
        # layout.addWidget(QLabel("Отображение ответов:"))
        # layout.addWidget(self.selectAnswers)
        #
        # # Кнопка экспорта
        # self.exportButton = QPushButton("Экспорт в Word")
        # layout.addWidget(self.exportButton)
        #
        # self.setLayout(layout)

    def setup_connections(self):
        self.chooseDirectoryButton.clicked.connect(self.choose_directory)
        self.exportButton.clicked.connect(self.export_to_word)

    def choose_directory(self):
        directory = QFileDialog.getExistingDirectory(
            self,
            "Выберите директорию",
            os.path.expanduser("~"),
            QFileDialog.Option.ShowDirsOnly
        )
        if directory:
            self.pathDirectory.setText(directory)

    def export_to_word(self):
        try:
            # Валидация данных
            if not self.pathDirectory.text():
                QMessageBox.warning(self, "Ошибка", "Выберите директорию для сохранения!")
                return

            if not self.fileName.text().strip():
                QMessageBox.warning(self, "Ошибка", "Введите название файла!")
                return

            # Формирование полного пути
            file_name = self.fileName.text().strip() + ".docx"
            full_path = os.path.join(self.pathDirectory.text(), file_name)

            # Создание документа
            doc = Document()
            self.configure_styles(doc)

            # Добавление контента в зависимости от выбранного варианта
            answer_option = self.selectAnswers.currentText()

            if answer_option == "показывать рядом с заданием":
                content = [f"{t}\nОтвет: {a.split(' ', 1)[1]}"
                           for t, a in zip(self.test_text, self.answers)]
            elif answer_option == "показывать в конце":
                content = self.test_text + ["\nОтветы:"] + self.answers
            else:
                content = self.test_text

            # Добавление текста в документ
            for paragraph in content:
                if paragraph.startswith("Ответ:"):
                    p = doc.add_paragraph()
                    run = p.add_run(paragraph)
                    run.italic = True
                else:
                    doc.add_paragraph(paragraph)

            # Сохранение документа
            doc.save(full_path)
            QMessageBox.information(self, "Успех", f"Файл сохранен:\n{full_path}")
            self.hide()
            del self

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка при экспорте: {str(e)}")

    def configure_styles(self, doc):
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        heading_style = doc.styles['Heading 1']
        heading_font = heading_style.font
        heading_font.bold = True
        heading_font.size = Pt(16)

