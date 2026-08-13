"""
Стандартные реализации Block.

TextBlock      — обычный текст
FormulaBlock   — LaTeX-формула, рендерится через matplotlib
ImageBlock     — растровое изображение (PIL.Image, bytes или путь)
CodeBlock      — листинг кода с моноширинным шрифтом
TableBlock     — табличные данные

Все Qt- и docx-зависимости импортируются ЛЕНИВО (внутри методов
render_qt/render_docx). Это позволяет тащить блоки в headless-окружения
(FastAPI, серверная сборка) без установленного PyQt6.

to_dict() — четвёртый метод полиморфного рендеринга. Возвращает
JSON-совместимый dict для веб-API. Бинарные данные кодируются в base64.
"""

from __future__ import annotations
import base64
import io
from pathlib import Path
from typing import Sequence, TYPE_CHECKING

from .content import Block

if TYPE_CHECKING:
    from PyQt6.QtWidgets import QWidget


# ---------- Вспомогательное ----------

def _b64(data: bytes) -> str:
    return base64.b64encode(data).decode("ascii")


# ---------- Блоки ----------

#: Кегли абзаца. Три, а не произвольное число: набор задан списком, потому
#: что «формат» здесь — стиль, а не вёрстка. Свободный размер в пунктах
#: означал бы, что автор верстает документ вручную, и два задания в одной
#: работе выглядели бы по-разному без всякой причины.
TEXT_SIZES = ("small", "normal", "large")

#: Кегль → пункты в .docx. В вебе размер задаётся классом, поэтому там
#: пунктов нет: у экрана своя типографика.
_DOCX_POINTS = {"small": 9, "normal": 11, "large": 14}


class TextBlock(Block):
    """
    Текстовый абзац с необязательным начертанием.

    Стиль появился ради редактора формата задания: условие бывает из
    нескольких абзацев, и заголовок раздела, само условие и примечание
    мелким шрифтом — это три разных абзаца, а не один слипшийся текст.

    Умолчания подобраны так, что `TextBlock("текст")` ведёт себя ровно как
    прежде: обычный кегль, без начертания. Это важно — конструктор зовут
    из полутора десятков мест, включая генераторы, которых стиль не
    касается вовсе.
    """

    def __init__(self, text: str, *, size: str = "normal",
                 bold: bool = False, italic: bool = False):
        self.text = text
        self.size = size if size in TEXT_SIZES else "normal"
        self.bold = bool(bold)
        self.italic = bool(italic)

    @property
    def styled(self) -> bool:
        """Отличается ли абзац от обычного — нужно и docx, и JSON."""
        return self.size != "normal" or self.bold or self.italic

    def render_qt(self, parent: "QWidget") -> "QWidget":
        from PyQt6.QtCore import Qt
        from PyQt6.QtWidgets import QLabel
        lbl = QLabel(self.text, parent)
        lbl.setWordWrap(True)
        lbl.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        if self.styled:
            font = lbl.font()
            font.setBold(self.bold)
            font.setItalic(self.italic)
            base = font.pointSize()
            if base > 0:
                font.setPointSize(
                    {"small": max(1, base - 2), "normal": base,
                     "large": base + 3}[self.size])
            lbl.setFont(font)
        return lbl

    def render_plain(self) -> str:
        return self.text

    def render_docx(self, doc) -> None:
        para = doc.add_paragraph()
        run = para.add_run(self.text)
        run.bold = self.bold
        run.italic = self.italic
        if self.size != "normal":
            from docx.shared import Pt
            run.font.size = Pt(_DOCX_POINTS[self.size])

    def to_dict(self) -> dict:
        out = {"type": "text", "content": self.text}
        # Поля добавляются только у оформленных абзацев: обычный блок
        # обязан сериализоваться байт в байт как раньше — его форму читают
        # десктоп, фронт и замороженные контрактные тесты.
        if self.styled:
            out["size"] = self.size
            out["bold"] = self.bold
            out["italic"] = self.italic
        return out


class FormulaBlock(Block):
    """LaTeX-формула."""

    def __init__(self, latex: str):
        self.latex = latex

    def render_qt(self, parent: "QWidget") -> "QWidget":
        from PyQt6.QtWidgets import QLabel
        from .rendering import latex_to_pixmap
        pix = latex_to_pixmap(self.latex)
        lbl = QLabel(parent)
        if pix is not None:
            lbl.setPixmap(pix)
        else:
            lbl.setText(f"${self.latex}$")
            lbl.setWordWrap(True)
        return lbl

    def render_plain(self) -> str:
        return f"${self.latex}$"

    def render_docx(self, doc) -> None:
        from .rendering import latex_to_docx_image
        latex_to_docx_image(doc, self.latex)

    def to_dict(self) -> dict:
        """
        Отдаём LaTeX-исходник. Картинку не отдаём — её рисует клиент.

        Раньше здесь готовился base64-PNG от matplotlib, и веб показывал
        его как `<img>`. Этап 7 плана (§10.2) называет это
        самостоятельным дефектом, и замеры это подтвердили: у типового
        задания матана PNG составлял **94 % ответа** (5432 байта из
        5760) и стоил около 40 мс на КАЖДУЮ формулу — в синхронном
        запросе, пока студент ждёт. Картинка при этом не выделяется, не
        копируется, не ищется по странице и не масштабируется вместе с
        текстом.

        Веб рисует формулы KaTeX по этому самому исходнику. Остальные
        два рендерера картинку берут не отсюда и не через этот метод:
        Qt — из `render_qt`, DOCX — из `render_docx`, каждый своим
        путём. То есть `image_b64` кормил ровно одного потребителя,
        который в нём больше не нуждается, и оставлять поле «на всякий
        случай» значило бы платить те же 94 % ни за что.
        """
        return {
            "type": "formula",
            "latex": self.latex,
        }


class ImageBlock(Block):
    """Изображение. Принимает PIL.Image, bytes или путь к файлу."""

    def __init__(self, image, caption: str = ""):
        self.image = image
        self.caption = caption

    def render_qt(self, parent: "QWidget") -> "QWidget":
        from PyQt6.QtCore import Qt
        from PyQt6.QtWidgets import QLabel
        from .rendering import pil_to_qpixmap
        pix = pil_to_qpixmap(self.image)
        lbl = QLabel(parent)
        if pix is not None:
            lbl.setPixmap(pix)
        else:
            lbl.setText(f"[{self.caption or 'изображение недоступно'}]")
        lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        return lbl

    def render_plain(self) -> str:
        return f"[изображение: {self.caption or 'без подписи'}]"

    def render_docx(self, doc) -> None:
        from PIL import Image as PILImage
        if isinstance(self.image, (str, Path)):
            doc.add_picture(str(self.image))
        elif isinstance(self.image, (bytes, bytearray)):
            doc.add_picture(io.BytesIO(self.image))
        elif isinstance(self.image, PILImage.Image):
            buf = io.BytesIO()
            self.image.save(buf, format="PNG")
            buf.seek(0)
            doc.add_picture(buf)
        else:
            doc.add_paragraph(f"[не удалось вставить изображение: {self.caption}]")

    def to_dict(self) -> dict:
        """
        PNG в base64 — здесь растр по делу, в отличие от формулы: у
        картинки исходника, из которого её можно нарисовать заново, нет.

        Три вида источника — те же, что у `render_docx`: путь, готовые
        байты, объект PIL. Раньше здесь вызывался `image_to_png_bytes`
        из `rendering`, которого в `rendering` нет — то есть ЛЮБОЕ
        задание с картинкой падало при сериализации с `ImportError`, а
        значит и в вебе не показывалось вовсе. Заметить это было негде:
        единственный тест на `ImageBlock.to_dict` лежит в файле-скрипте,
        который `unittest discover` не собирает.

        Не удалось — `None`, как обещает контракт `Block.to_dict`: без
        картинки задание всё ещё читается, а исключение отсюда уронило
        бы весь ответ.
        """
        png = self._png_bytes()
        return {
            "type": "image",
            "image_b64": _b64(png) if png is not None else None,
            "caption": self.caption,
        }

    def _png_bytes(self) -> bytes | None:
        try:
            if isinstance(self.image, (bytes, bytearray)):
                return bytes(self.image)
            if isinstance(self.image, (str, Path)):
                return Path(self.image).read_bytes()
            buf = io.BytesIO()
            self.image.save(buf, format="PNG")
            return buf.getvalue()
        except Exception:                                  # noqa: BLE001
            return None


class CodeBlock(Block):
    """Листинг кода."""

    def __init__(self, code: str, language: str = "text"):
        self.code = code
        self.language = language

    def render_qt(self, parent: "QWidget") -> "QWidget":
        from PyQt6.QtGui import QFont
        from PyQt6.QtWidgets import QPlainTextEdit
        edit = QPlainTextEdit(parent)
        edit.setPlainText(self.code)
        edit.setReadOnly(True)
        font = QFont("Consolas")
        font.setStyleHint(QFont.StyleHint.Monospace)
        font.setPointSize(10)
        edit.setFont(font)
        return edit

    def render_plain(self) -> str:
        return f"```{self.language}\n{self.code}\n```"

    def render_docx(self, doc) -> None:
        p = doc.add_paragraph()
        run = p.add_run(self.code)
        run.font.name = "Consolas"

    def to_dict(self) -> dict:
        return {
            "type": "code",
            "code": self.code,
            "language": self.language,
        }


class TableBlock(Block):
    """Таблица."""

    def __init__(
        self,
        rows: Sequence[Sequence[str]],
        header: Sequence[str] | None = None,
    ):
        self.rows = [list(r) for r in rows]
        self.header = list(header) if header else None

    def render_qt(self, parent: "QWidget") -> "QWidget":
        from PyQt6.QtWidgets import QTableWidget, QTableWidgetItem
        cols = len(self.header) if self.header else (len(self.rows[0]) if self.rows else 0)
        tbl = QTableWidget(len(self.rows), cols, parent)
        if self.header:
            tbl.setHorizontalHeaderLabels(self.header)
        for r, row in enumerate(self.rows):
            for c, val in enumerate(row):
                tbl.setItem(r, c, QTableWidgetItem(str(val)))
        tbl.resizeColumnsToContents()
        return tbl

    def render_plain(self) -> str:
        out = []
        if self.header:
            out.append(" | ".join(self.header))
            out.append("-" * len(out[0]))
        for row in self.rows:
            out.append(" | ".join(str(c) for c in row))
        return "\n".join(out)

    def render_docx(self, doc) -> None:
        cols = len(self.header) if self.header else (len(self.rows[0]) if self.rows else 0)
        if cols == 0:
            return
        tbl = doc.add_table(rows=(1 if self.header else 0) + len(self.rows), cols=cols)
        tbl.style = "Light Grid Accent 1"
        ofs = 0
        if self.header:
            for c, h in enumerate(self.header):
                tbl.rows[0].cells[c].text = h
            ofs = 1
        for r, row in enumerate(self.rows):
            for c, val in enumerate(row):
                tbl.rows[r + ofs].cells[c].text = str(val)

    def to_dict(self) -> dict:
        return {
            "type": "table",
            "rows": [[str(c) for c in row] for row in self.rows],
            "header": list(self.header) if self.header else None,
        }


# ---------- Обратный разбор ----------
#
# to_dict() у блоков был односторонним: веб-сериализация отдаёт словарь,
# и обратно его никто не собирал — фронту хватало «type» для выбора
# компонента. Общей интерактивной сессии этого мало: её снимок состояния
# содержит условие вопроса, и при переезде между процессами условие надо
# восстановить, а не показать пользователю другое.

def block_from_dict(data: dict) -> Block:
    """
    Собрать блок из словаря, выданного `to_dict()`.

    Обратимы блоки, которые целиком состоят из данных. Картинка
    восстанавливается из base64-PNG: исходный PIL.Image или путь не
    сохраняются, но показывается ровно то же самое.

    Неизвестный тип — ValueError, а не «молча TextBlock»: подмена блока
    заглушкой выглядит как испорченное задание, и искать причину придётся
    в отрендеренном виде, где её уже не видно.
    """
    if not isinstance(data, dict):
        raise ValueError(f"Блок должен быть словарём, получено {type(data).__name__}")
    kind = data.get("type")
    builder = _BLOCK_BUILDERS.get(kind)
    if builder is None:
        raise ValueError(f"Неизвестный тип блока: {kind!r}")
    return builder(data)


def blocks_from_dicts(items) -> list:
    """Собрать список блоков. Пустой вход — пустой список."""
    return [block_from_dict(item) for item in (items or [])]


def _image_from_dict(data: dict) -> Block:
    raw = data.get("image_b64")
    payload = base64.b64decode(raw) if raw else b""
    return ImageBlock(payload, caption=data.get("caption", ""))


_BLOCK_BUILDERS = {
    "text": lambda d: TextBlock(d.get("content", "")),
    "formula": lambda d: FormulaBlock(d.get("latex", "")),
    "code": lambda d: CodeBlock(d.get("code", ""), d.get("language", "text")),
    "table": lambda d: TableBlock(d.get("rows") or [], d.get("header")),
    "image": _image_from_dict,
}
