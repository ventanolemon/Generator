"""
Экспорт заданий в .docx — единственная реализация для всех предметов.

Не знает о предметах. Работает с list[StaticTask] и зовёт у каждого
блока render_docx — как и view-слой.
"""

from __future__ import annotations
from typing import List

from docx import Document
from docx.shared import Pt

from core import StaticTask, TextBlock


def export_tasks_to_docx(
    tasks: List[StaticTask],
    path: str,
    title: str = "Задания",
    with_answers: bool = True,
) -> None:
    """
    Каждое задание — пронумерованный блок. Если with_answers, в конце
    отдельный раздел 'Ответы' с такой же нумерацией.
    """
    doc = Document()
    h = doc.add_heading(title, level=0)

    for i, task in enumerate(tasks, 1):
        doc.add_heading(f"Задание {i}", level=2)
        for block in task.statement:
            block.render_docx(doc)

    if with_answers and tasks:
        doc.add_page_break()
        doc.add_heading("Ответы", level=1)
        for i, task in enumerate(tasks, 1):
            doc.add_heading(f"Ответ {i}", level=2)
            for block in task.answer:
                block.render_docx(doc)

    doc.save(path)


def export_test_to_docx(
    variants: List[StaticTask],
    path: str,
    title: str = "Тест",
    with_answers: bool = False,
) -> None:
    """
    Каждый вариант — отдельный раздел. Если with_answers, в конце —
    эталон ответов всех вариантов.
    """
    doc = Document()
    doc.add_heading(title, level=0)

    for i, variant in enumerate(variants, 1):
        doc.add_heading(f"Вариант {i}", level=1)
        for block in variant.statement:
            block.render_docx(doc)
        if i < len(variants):
            doc.add_page_break()

    if with_answers:
        doc.add_page_break()
        doc.add_heading("Эталон ответов", level=1)
        for i, variant in enumerate(variants, 1):
            doc.add_heading(f"Вариант {i}", level=2)
            for block in variant.answer:
                block.render_docx(doc)

    doc.save(path)
